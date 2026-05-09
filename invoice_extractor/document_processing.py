from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path


class DocumentExtractionError(RuntimeError):
    """Raised when a document cannot be converted into text."""


@dataclass
class ExtractionArtifact:
    source_name: str
    source_type: str
    text: str
    page_count: int | None = None
    used_ocr: bool = False
    warnings: list[str] = field(default_factory=list)


def extract_text_from_upload(
    filename: str,
    content: bytes,
    ocr_language: str,
    min_direct_pdf_chars: int,
) -> ExtractionArtifact:
    suffix = Path(filename or "upload").suffix.lower()

    if suffix == ".pdf":
        text, page_count, used_ocr, warnings = _extract_text_from_pdf(
            content=content,
            ocr_language=ocr_language,
            min_direct_pdf_chars=min_direct_pdf_chars,
        )
        return ExtractionArtifact(
            source_name=filename,
            source_type="PDF",
            text=text,
            page_count=page_count,
            used_ocr=used_ocr,
            warnings=warnings,
        )

    if suffix == ".docx":
        return ExtractionArtifact(
            source_name=filename,
            source_type="DOCX",
            text=_extract_text_from_docx(content),
            warnings=[],
        )

    if suffix in {".png", ".jpg", ".jpeg"}:
        return ExtractionArtifact(
            source_name=filename,
            source_type="Image",
            text=_extract_text_from_image(content, ocr_language),
            used_ocr=True,
            warnings=[],
        )

    raise DocumentExtractionError("Unsupported file type. Use PDF, DOCX, PNG, JPG or JPEG.")


def _extract_text_from_pdf(
    content: bytes,
    ocr_language: str,
    min_direct_pdf_chars: int,
) -> tuple[str, int, bool, list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError("The package 'pypdf' is required to read PDF files.") from exc

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError(f"Unable to open the PDF file: {exc}") from exc

    page_texts: list[str] = []
    for page in reader.pages:
        page_texts.append((page.extract_text() or "").strip())

    combined_text = _clean_extracted_text("\n\n".join(page_texts))
    page_count = len(reader.pages)
    warnings: list[str] = []
    used_ocr = False

    if len(combined_text) >= min_direct_pdf_chars:
        return combined_text, page_count, used_ocr, warnings

    ocr_text, ocr_warning = _extract_text_from_pdf_with_ocr(content, ocr_language)
    if ocr_warning:
        warnings.append(ocr_warning)

    if ocr_text:
        used_ocr = True
        combined_text = _clean_extracted_text("\n\n".join(filter(None, [combined_text, ocr_text])))

    if not combined_text:
        details = (
            "No readable text was found in the PDF. "
            "Use a text-based PDF or install OCR dependencies for scanned documents."
        )
        if ocr_warning:
            details = f"{details} Details: {ocr_warning}"
        raise DocumentExtractionError(details)

    if used_ocr:
        warnings.append("OCR fallback was used because the PDF text layer was limited.")
    elif len(combined_text) < min_direct_pdf_chars:
        warnings.append("The PDF contains little extractable text. Results may be incomplete.")

    return combined_text, page_count, used_ocr, warnings


def _extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("The package 'python-docx' is required to read DOCX files.") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError(f"Unable to open the DOCX file: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = _clean_extracted_text("\n".join(lines))
    if not text:
        raise DocumentExtractionError("No readable text was found in the DOCX file.")
    return text


def _extract_text_from_image(content: bytes, ocr_language: str) -> str:
    image = _load_image(content)
    text = _run_ocr(image, ocr_language)
    cleaned = _clean_extracted_text(text)
    if not cleaned:
        raise DocumentExtractionError("OCR did not detect readable text in the image.")
    return cleaned


def _extract_text_from_pdf_with_ocr(content: bytes, ocr_language: str) -> tuple[str, str | None]:
    try:
        import fitz
    except ImportError:
        return "", "PyMuPDF is not installed, so scanned PDF OCR is unavailable."

    try:
        document = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        return "", f"Unable to render the PDF for OCR: {exc}"

    page_texts: list[str] = []
    try:
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = _load_image(pixmap.tobytes("png"))
            page_text = _run_ocr(image, ocr_language)
            if page_text.strip():
                page_texts.append(page_text.strip())
    except DocumentExtractionError as exc:
        return "", str(exc)
    finally:
        document.close()

    return _clean_extracted_text("\n\n".join(page_texts)), None


def _load_image(content: bytes):
    try:
        from PIL import Image
    except ImportError as exc:
        raise DocumentExtractionError("The package 'Pillow' is required for image handling.") from exc

    try:
        return Image.open(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError(f"Unable to open the image file: {exc}") from exc


def _run_ocr(image, ocr_language: str) -> str:
    try:
        import pytesseract
        from PIL import ImageFilter, ImageOps
    except ImportError as exc:
        raise DocumentExtractionError(
            "Image OCR requires 'pytesseract' and 'Pillow'. Install the runtime dependencies first."
        ) from exc

    prepared = ImageOps.grayscale(image.convert("RGB"))
    prepared = ImageOps.autocontrast(prepared.filter(ImageFilter.SHARPEN))

    try:
        pytesseract.get_tesseract_version()
        return pytesseract.image_to_string(prepared, lang=ocr_language)
    except Exception as exc:
        raise DocumentExtractionError(
            "Tesseract OCR is not available. Install Tesseract locally or add it in deployment."
        ) from exc


def _clean_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    compact_lines = [line for line in lines if line]
    return "\n".join(compact_lines).strip()
